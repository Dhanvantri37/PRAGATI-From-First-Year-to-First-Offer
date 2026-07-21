import docx
doc = docx.Document(r"d:\Hacker\Aptitude Sheet (1).docx")
for idx in range(2824, 2860):
    if idx < len(doc.paragraphs):
        print(f"P{idx}: {repr(doc.paragraphs[idx].text)}")
